#!c:\python36\python.exe
'''
Created on Dec 12, 2018

'''
import time
import traceback
import os
from Logger import Logger
from faxJob import faxJob
import Faxer
from threading import Thread
import docx
import re
import subprocess
from shutil import move
import Needles
from fileinput import filename
import random
from xmlgen import generateXml

class JobHandler(object):
    running=True
    idle=10
    watchFile=""
    failedFolder="failed"
    delOrMoveCompleted=""
    moveFolder="sent"
    storeToDW=True
    libreEXE=""
    processes=[]
    faxableFiles=['DOC','DOCX','PDF','PNG','JPG','BMP','JPEG','TIF','TIFF','TXT']

    def __init__(self, idle,  watchFile, delOrMoveCompleted, moveFolder, storeToDW, libreEXE, failedFolder):
        self.idle=idle
        self.watchFile = watchFile
        self.delOrMoveCompleted=delOrMoveCompleted
        self.moveFolder=moveFolder
        self.storeToDW=storeToDW
        self.libreEXE=libreEXE
        self.failedFolder=failedFolder
        
    def run(self):
        print("JobHandler run")
        try:
            while(self.running):
                self.handleJobs()
                time.sleep(self.idle)
        except:
            Logger.writeAndPrintLine("An unexpected error occurred in JobHandler, halting: "+traceback.format_exc(),3)  
  
    def handleJobs(self):
        for job in faxJob.faxJobs:
            if(job.status=="MATCHED"):
                #jobThread=Thread(target=self.runJob, args=(job,))
                #jobThread.start()
                Logger.writeAndPrintLine("Starting job for "+job.uid,5)
                self.runJob(job)
            if(job.status=="COMPLETE"):
                Logger.writeAndPrintLine("Fax job for "+job.uid+" complete.",1)
                faxJob.faxJobs.remove(job)
    
    def runJob(self, job):
        job.status="RUNNING"
        try:
            if(self.pullFaxFields(job)):
                job.status="PARSED"
            else:
                Logger.writeAndPrintLine("Could not find sufficient faxing/storing information for job "+job.uid+'.',3)
                self.sendErrorMessage(job, "Could not find sufficient faxing/storing information for job "+job.uid+' Documents not faxed, see '+self.failedFolder)
                job.status="ERRORED"
        except:
            Logger.writeAndPrintLine("Failed to pull hidden fields from auth, document may be malformed? Does fax number exist for provider? "
                                     +traceback.format_exc(),3) 
            job.status="ERRORED"
            return
        
        doStore=not self.pruneCoverSheet(job)
        if(doStore):
            if(job.status!="ERRORED"):
                if(self.storeToDW and job.status!="ERRORED"):
                    result=self.storeToDocuware(job)
                    if(result==None):
                        Logger.writeAndPrintLine("Successfully stored "+job.uid+" to Docuware.",1) 
                        job.status="STORED"
                    else:
                        Logger.writeAndPrintLine("Error storing document(s) to Docuware: "+str(result),3) 
                        self.sendErrorMessage(job, "Error storing documents for job "+job.uid+' Documents not faxed, see '+self.watchFile)
                        job.status="ERRORED"
        
        if(job.status!="ERRORED"):
            try:
                Logger.writeAndPrintLine("Attempting to convert document(s) to PDF. ",5) 
                self.convertDocsToPDF(job)
            except:
                Logger.writeAndPrintLine("Failed to convert document(s) to PDF. "+traceback.format_exc(),3) 
                job.status="ERRORED"
                return
        
        if(job.status!="ERRORED"):
            try:
                self.sendFax(job)
                Logger.writeAndPrintLine("Fax sent successfully to "+job.destFax+". "+job.uid,1) 
                job.status="SENT"
            except:
                Logger.writeAndPrintLine("Error sending fax through faxfinder - network error? "+traceback.format_exc(),3) 
                job.status="ERRORED"
                return
            
        if(job.status!="ERRORED"):
            try:
                self.cleanupJob(job)
                job.status="COMPLETE"
            except:
                Logger.writeAndPrintLine("Error cleaning up job - access to files or move dir? "+traceback.format_exc(),3) 
                job.status="ERRORED"
                return
        if(job.status=="ERRORED"):
            if(self.moveFailedJob(job)):
                Logger.writeAndPrintLine("Moved failed job "+job.uid+" to "+self.failedFolder,3) 
            job.status="DEAD"

    def pullFaxFields(self, job):
        for file in job.files:
            if('.DOC' in file.upper()):
                print('trying '+file)
                try:
                    #https://automatetheboringstuff.com/chapter13/
                    doc = docx.Document(file)
                    fullText = []
                    for para in doc.paragraphs:
                        fullText.append(para.text)
                    fulltext='\n'.join(fullText)
                    #print(fulltext)
                    
                    job.dwCasenum=re.search('dwCasenum=(.+)',fulltext,0).group(1)
                    job.dwName=re.search('dwName=(.+)',fulltext,0).group(1)
                    job.dwProvider=re.search('dwProvider=(.+)',fulltext,0).group(1)
                    job.dwCategory=re.search('dwCategory=(.+)',fulltext,0).group(1)
                    job.dwDocumentType=re.search('dwDocumentType=(.+)',fulltext,0).group(1)
                    
                    job.comment=re.search('ffComment=(.+)',fulltext,0).group(1)
                    job.destName=re.search('ffDestName=(.+)',fulltext,0).group(1)
                    job.destOrg=re.search('ffDestOrg=(.+)',fulltext,0).group(1)
                    job.destFax=re.search('ffDestFax=(.+)',fulltext,0).group(1)
                    job.cliName=re.search('ffCliName=(.+)',fulltext,0).group(1)
                    job.casenum=re.search('ffCasenum=(.+)',fulltext,0).group(1)  
                    job.staffEmail=re.search('staffEmail=(.+)',fulltext,0).group(1)  

                    #generateXml(job.cliName, job.casenum, job.comment, job.destOrg, job.destFax, job.staffEmail, job.destName)

                    if(job.destFax!=""):
                        break
                except: 
                    Logger.writeAndPrintLine("Error pulling info for file "+file+" "+traceback.format_exc(),3) 
                    #print("Error parsing file "+file+": "+traceback.format_exc())
                    None
        if(job.destFax==""):
            return False
        else:
            job.destFax=job.destFax.replace('(','').replace(')','').replace('-','').replace(' ','')
            job.destFax='1'+job.destFax
            return True
        print(job.destFax)

    
    def pruneCoverSheet(self, job):
        for file in job.files:
            if('.DOC' in file.upper()):
                try:
                    doc = docx.Document(file)
                    fullText = []
                    for para in doc.paragraphs:
                        fullText.append(para.text)
                    fulltext='\n'.join(fullText)
                    if("coverPage=TRUE" in fulltext):
                        Logger.writeAndPrintLine("Removing file "+file,5) 
                        job.files.remove(file)
                        if(self.delOrMoveCompleted=="DELETE"):
                            os.unlink(file)
                        else:
                            move(file, self.moveFolder)
                        Logger.writeAndPrintLine("Cover sheet removed from "+job.uid,1) 
                        return True
                except:
                    None
        return False
                    
    
    def convertDocsToPDF(self, job):
        i=-1
        while(i<len(job.files)-1):
            i=i+1
            if(job.files[i].split('.')[-1].upper() in self.faxableFiles and ".PDF" not in job.files[i].upper()):
                pdfFile="".join(job.files[i].split('.')[:-1])+'.pdf'
                if(os.path.exists(pdfFile)):
                    os.unlink(pdfFile)
                    
                #there's a bug where libre convert will hang without using a random temp folder https://bugs.documentfoundation.org/show_bug.cgi?id=113259
                subprocess.Popen([self.libreEXE,"--headless","--convert-to","pdf",job.files[i],"--outdir",self.watchFile], stdout=subprocess.PIPE)
                #output=p1.communicate()[0]
                #Logger.writeAndPrintLine(output.decode('ascii'),5)
                #wait for libre to create the PDF file. 
                while(True):
                    if(os.path.exists(pdfFile)):
                        Logger.writeAndPrintLine("Converted "+job.files[i]+" to pdf.",1)
                        break
                    time.sleep(1)
                #wait for libre to release the word doc. 
                while(True):
                    try:
                        os.unlink(job.files[i])
                        Logger.writeAndPrintLine("Cleaned up old doc "+job.files[i],5)
                        job.files[i]=pdfFile
                        break
                    except:
                        time.sleep(1)
                 
        
    def sendFax(self, job):
        #generateXml(job.cliName, job.casenum, job.comment, job.destOrg, job.destFax, job.staffEmail, job.destName, job.files)
        Faxer.sendFax(job.destOrg, job.destFax, job.cliName, job.casenum, job.files, job.staffEmail, job.comment, job.destName)
        

    def cleanupJob(self, job):
        if(self.delOrMoveCompleted=="DELETE"):
            for file in job.files:
                os.unlink(file)
        else:
            for file in job.files:
                move(file, self.moveFolder)
            
    def storeToDocuware(self, job):
        dwStorer="DWStorer\DWStorer.exe"
        params=[]
        params.append(dwStorer)
        params.append('-c')
        params.append('-iCASE_ID:"'+job.dwCasenum+'"')
        params.append('-iLAST_NAME:"'+job.dwName+'"')
        params.append('-iPROVIDER:"'+job.dwProvider+'"')
        params.append('-iCATEGORY:"'+job.dwCategory+'"')
        params.append('-iDOCUMENT_TYPE:"'+job.dwDocumentType+'"')
        params.append('-iSTATUS:"NEW"')
        params.append('-iSTOREDBY:"AutoFaxer"')
        for file in job.files:
            params.append('-f"'+file+'"')
        out=subprocess.check_output(params)
        out=out.decode('ascii')
        if("success" in out):
            return None
        else:
            return out
        
    def sendErrorMessage(self, job, body):
        recipients=job.staffcode+"; MNOW"
        Needles.sendMessage(job.staffcode, recipients, body, None, None, None)
        Needles.sendMessage("MNOW", recipients, body, None, None, None)
        Logger.writeAndPrintLine("Sent error messages successfully regarding job "+job.uid,2) 
        
    def moveFailedJob(self, job):
        for file in job.files:
            filename=file.split('\\')[-1]
            newFullpath=self.failedFolder+'\\'+filename
            try:
                os.rename(file, newFullpath)
            except: 
                Logger.writeAndPrintLine("Failed to move file "+file+": "+traceback.format_exc(),3) 
                return False
        return True
            
        